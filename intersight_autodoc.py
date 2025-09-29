import os
import json
import requests
import pandas as pd
import yaml
import sys
from dotenv import load_dotenv
from intersight_auth import IntersightAuth
from docx import Document
from docx.shared import Pt
from requests import Session

load_dotenv(dotenv_path=".env")

# Validate configuration files
try:
    with open('config.yaml', 'r') as config_file:
        config = yaml.safe_load(config_file)
        CUSTOMER_INFO = config['CUSTOMER_INFO']
except FileNotFoundError:
    raise FileNotFoundError("The 'config.yaml' file is missing. Please provide it.")
except yaml.YAMLError as e:
    raise ValueError(f"Error parsing 'config.yaml': {e}")

try:
    with open('operations.yaml', 'r') as file:
        operations_data = yaml.safe_load(file)
        OPERATIONS = operations_data['OPERATIONS']
except FileNotFoundError:
    raise FileNotFoundError("The 'operations.yaml' file is missing. Please provide it.")
except yaml.YAMLError as e:
    raise ValueError(f"Error parsing 'operations.yaml': {e}")

# Access configuration values
OUTPUT_DIRECTORY = config['directories']['output'] # Directory containing JSON files
FLATTENED_OUTPUT_DIRECTORY = config['directories']['flattened_output'] # Directory to save flattened JSON files
EXCEL_OUTPUT_DIRECTORY = config['directories']['excel_output'] # Directory to save Excel files
AUTODOC_TEMPLATE_PATH = config['autodoc_template_path'] # Path to the Word document template
AUTODOC_COMPLETED_PATH = config['autodoc_completed_path'] # Path to the completed Word document 
SECRET_KEY_PATH = os.getenv("SECRET_KEY_PATH") # Intersight API Secret Key
API_KEY_ID = os.getenv("API_KEY_ID") # Intersight API Key
BURL = config['base_url']  # Intersight REST API Base URL


# Ensure the directories exist
def ensure_directories_exist():
    """
    Ensures that all required directories exist. If a directory does not exist, it is created.

    This function is typically called at the start of the program to prepare the environment
    for file operations such as saving JSON or Excel files.
    """
    os.makedirs(OUTPUT_DIRECTORY, exist_ok=True)
    os.makedirs(FLATTENED_OUTPUT_DIRECTORY, exist_ok=True)
    os.makedirs(EXCEL_OUTPUT_DIRECTORY, exist_ok=True)

def create_auth_object():
    """
    Creates and returns an authentication object for Intersight API.

    This function initializes an `IntersightAuth` object using the secret key file
    and API key ID defined in the environment variables.

    Returns:
        IntersightAuth: An authentication object to be used for API requests.
    """

    with open(SECRET_KEY_PATH, 'r') as key_file:
        my_secret_key = key_file.read()

    session = Session()
    session.auth = IntersightAuth(
        api_key_id=API_KEY_ID,
        secret_key_string=my_secret_key
    )
    return session.auth
    

def get_nested_value(data, keys):
    """
    Retrieves a nested value from a dictionary or list based on a sequence of keys.

    Args:
        data (dict or list): The data structure (dictionary or list) to search.
        keys (list): A list of keys representing the path to the desired value.

    Returns:
        Any: The value found at the specified path, or None if the path does not exist.
    """
    for key in keys:
        if isinstance(data, list):
            # Iterate over the list and extract all matching values
            data = [get_nested_value(item, keys[1:]) for item in data]
            data = [item for item in data if item is not None]
            if len(data) == 1:
                data = data[0]  # Flatten single-item lists
        elif isinstance(data, dict):
            data = data.get(key)
        else:
            return None
    return data

def filter_json(response_json, filter_keys):
    """
    Filters a JSON response to include only specified keys, supporting nested key paths.

    Args:
        response_json (dict): The JSON response to filter, expected to contain a 'Results' key.
        filter_keys (list): A list of keys (including nested keys separated by '.') to extract.

    Returns:
        list: A list of dictionaries containing the filtered key-value pairs.

    Notes:
        - Nested keys are resolved using the `get_nested_value` function.
        - Lists of dictionaries are flattened into readable JSON strings.
        - If a key does not exist, its value in the filtered result will be None.
    """
    def get_nested_value(data, keys):
        for key in keys:
            if isinstance(data, list):
                # If the current data is a list, iterate over each item
                data = [get_nested_value(item, [key]) for item in data]
                data = [item for item in data if item is not None]  # Remove None values
                if len(data) == 1:
                    data = data[0]  # Flatten single-item lists
            elif isinstance(data, dict):
                # If the current data is a dictionary, get the value for the key
                data = data.get(key)
            else:
                return None  # Key does not exist
        return data

    filtered_results = []
    for item in response_json.get('Results', []):  # Ensure 'Results' key exists
        filtered_item = {}
        for key in filter_keys:
            keys = key.split('.')
            value = get_nested_value(item, keys)
            # Flatten lists of dictionaries into readable strings
            if isinstance(value, list) and all(isinstance(v, dict) for v in value):
                value = [json.dumps(v) for v in value]
            filtered_item[key] = value
        filtered_results.append(filtered_item)
    return filtered_results



def main():
    ensure_directories_exist()
    auth = create_auth_object()
    
    # Load the Word document template
    doc = Document(AUTODOC_TEMPLATE_PATH)

    # Update document with placeholders from CUSTOMER_INFO in the config.yaml file.
    for config in CUSTOMER_INFO:
        for key, value in config.items():
            placeholder = f'{{{{{key}}}}}'  # Placeholders in the document are expected to be in the format {{key}}
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))


    # Process each operation defined in operations.yaml
    for operation in OPERATIONS:
        if operation['request_process']:
            response = None

            # GET
            if operation['request_method'] == "GET":
                response = requests.get(
                    BURL + operation['resource_path'] + operation['select'] + operation['expand'],
                    auth=auth
                )
            
            # Parse the JSON response
            response_json = response.json()

            # Check if the response contains an empty "Results" list
            if not response_json or ('Results' in response_json and not response_json['Results']):
                print(f"No data returned for operation: {operation['resource_path']}. Skipping...")
                continue

            # Save the JSON response to a separate file named after the resource_path
            resource_name = operation['resource_path'].replace('/', '_')
            output_file_path = os.path.join(OUTPUT_DIRECTORY, f'{resource_name}.json')
            with open(output_file_path, 'w') as output_file:
                json.dump(response_json, output_file, indent=4)

            # Filter JSON response to only include selected keys if filter exists
            if 'filter' in operation:
                filter_keys = [key.strip() for key in operation['filter'].split(',')]
                filtered_json = filter_json(response_json, filter_keys)
                filtered_output_file_path = os.path.join(FLATTENED_OUTPUT_DIRECTORY, f'filtered_{resource_name}.json')
                with open(filtered_output_file_path, 'w') as filtered_output_file:
                    json.dump(filtered_json, filtered_output_file, indent=4)

                # Convert filtered JSON response to Excel table
                df = pd.json_normalize(filtered_json)
                excel_output_file_path = os.path.join(EXCEL_OUTPUT_DIRECTORY, f'{resource_name}.xlsx')

                # Rename column names
                # Splitting of column_names
                column_names = operation.get('column_names', df.columns).split(',')

                # Ensure that splitting results in correct names
                column_names = [name.strip() for name in column_names]

                # Check if the length of column_names matches the number of columns in the DataFrame
                if len(column_names) == len(df.columns):
                    df.rename(columns=dict(zip(df.columns, column_names)), inplace=True)
                # else:
                #     raise ValueError("Mismatch in the number of column names and DataFrame columns.")
                excel_output_file_path = os.path.join(EXCEL_OUTPUT_DIRECTORY, f'{resource_name}.xlsx')
                df.to_excel(excel_output_file_path, index=False)

                placeholder = operation.get('placeholder', None)
                placeholder_found = False  # Flag to track if a placeholder is found

                if placeholder:
                    for paragraph in doc.paragraphs:
                        if placeholder in paragraph.text:
                            placeholder_found = True
                            # Replace the placeholder with an empty string
                            paragraph.text = paragraph.text.replace(placeholder, "")

                            # Add a heading for the table
                            table_name = operation.get('table_name', resource_name)
                            heading = paragraph.insert_paragraph_before(table_name)
                            heading.style = 'Heading 5-No Numbers'
                            for run in heading.runs:
                                run.font.size = Pt(10)

                            # Create the table
                            table = doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])
                            table.style = 'Scroll Table Normal'

                            # Move the table directly after the heading
                            heading._element.addnext(table._element)

                            # Add column headers
                            for j, col in enumerate(df.columns):
                                cell = table.cell(0, j)
                                cell.text = col
                                cell.paragraphs[0].runs[0].font.size = Pt(8)

                            # Add data rows
                            for i in range(df.shape[0]):
                                for j in range(df.shape[1]):
                                    cell = table.cell(i + 1, j)
                                    value = df.iat[i, j]
                                    if isinstance(value, list):  # Check if the value is a list
                                        value = '\n'.join(map(str, value))  # Convert all elements to strings and join with newlines
                                    elif isinstance(value, str) and ',' in value:
                                        value = '\n'.join(value.split(', '))  # Split string by commas and join with newlines
                                    cell.text = str(value)
                                    cell.paragraphs[0].runs[0].font.size = Pt(6.5)

                            # Exit the loop after processing the placeholder
                            break
                # If no placeholder is found, append the table at the end
                if not placeholder_found:
                    # print(f"No placeholder found for operation: {operation}")  # Debugging output
                    table_name = operation.get('table_name', resource_name)
                    # print(f"Appending table with name: {table_name}")  # Debugging output

                    # Add a heading for the table
                    heading = doc.add_heading(table_name, level=1)
                    heading.style = 'Heading 5-No Numbers'
                    for run in heading.runs:
                        run.font.size = Pt(10)

                    # Create the table
                    table = doc.add_table(rows=(df.shape[0] + 1), cols=df.shape[1])
                    table.style = 'Scroll Table Normal'

                    # Add column headers
                    for j, col in enumerate(df.columns):
                        cell = table.cell(0, j)
                        cell.text = col
                        cell.paragraphs[0].runs[0].font.size = Pt(8)

                    # Add data rows
                    for i in range(df.shape[0]):
                        for j in range(df.shape[1]):
                            cell = table.cell(i + 1, j)
                            value = df.iat[i, j]
                            if isinstance(value, list):  # Check if the value is a list
                                value = '\n'.join(map(str, value))  # Convert all elements to strings and join with newlines
                            elif isinstance(value, str) and ',' in value:
                                value = '\n'.join(value.split(', '))  # Split string by commas and join with newlines
                            cell.text = str(value)
                            cell.paragraphs[0].runs[0].font.size = Pt(6.5)
    # Save the updated Word document
    doc.save(AUTODOC_COMPLETED_PATH)

if __name__ == "__main__":
    main()
